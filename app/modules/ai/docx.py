from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.shared import Pt

from app.modules.ai.schemas import ExecutiveSummaryResponse


def render_executive_summary_docx(summary: ExecutiveSummaryResponse) -> bytes:
    doc = Document()

    doc.add_heading("Executive Summary", level=1)

    if summary.assessment_id:
        doc.add_paragraph(f"Assessment ID: {summary.assessment_id}")
    if summary.generated_at:
        doc.add_paragraph(f"Generated: {summary.generated_at}")
    if summary.maturity_band or summary.average_score is not None or summary.maturity_level is not None:
        meta = []
        if summary.maturity_band:
            meta.append(f"Maturity: {summary.maturity_band}")
        if summary.average_score is not None:
            meta.append(f"Avg Score: {summary.average_score}/100")
        if summary.maturity_level is not None:
            label = summary.maturity_label or ""
            meta.append(f"Level {summary.maturity_level}{f' ({label})' if label else ''}")
        doc.add_paragraph(" · ".join(meta))

    doc.add_paragraph(summary.summary_text or "Summary unavailable.")

    doc.add_heading("Key Findings", level=2)
    findings = summary.key_findings or []
    if findings:
        for item in findings:
            doc.add_paragraph(item, style="List Bullet")
    else:
        doc.add_paragraph("No findings generated yet.")

    doc.add_heading("High-Level Recommendations", level=2)
    recommendations = summary.high_level_recommendations or []
    if recommendations:
        for item in recommendations:
            doc.add_paragraph(item, style="List Bullet")
    else:
        doc.add_paragraph("No recommendations generated yet.")

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(11)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
